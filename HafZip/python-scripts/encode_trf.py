# encoding=utf-8

import docx
import sys


def transfer(doc_name):
    text_name = "D:\\USBei\\C++Projects\\Zip_Tool\\Zip_Tool\\1_doc.txt"
    f_txt = open(text_name, mode='w')

    # 获取文档对象

    doc_name = str(doc_name)
    file = docx.Document(doc_name)

    print("段落数:" + str(len(file.paragraphs)))  # 段落数为...，每个回车隔离一段

    # 输出每一段的内容

    for para in file.paragraphs:
        print(para.text)
        f_txt.write(para.text + '\n')

    f_txt.close()

transfer(sys.argv[1])
